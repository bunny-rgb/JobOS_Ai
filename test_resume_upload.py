#!/usr/bin/env python3
"""
JobOS AI Phase 3 - Resume Upload Endpoint Test Suite
Tests the FIXED /api/profile/resume-upload endpoint with pdf-parse v2 PDFParse class
"""

import requests
import time
import os
import sys

# Base URL from environment
BASE_URL = "https://job-os-app.preview.emergentagent.com/api"

# Test data
timestamp = int(time.time())
TEST_EMAIL = f"resume_test_{timestamp}@jobos.ai"
TEST_PASSWORD = "SecurePass123!"
TEST_NAME = "John Doe - Senior QA Automation Engineer"

# Global token
token = None

def print_test(name):
    print(f"\n{'='*80}")
    print(f"TEST: {name}")
    print('='*80)

def print_success(msg):
    print(f"✅ SUCCESS: {msg}")

def print_error(msg):
    print(f"❌ ERROR: {msg}")

def print_info(msg):
    print(f"ℹ️  INFO: {msg}")

# ============================================================================
# 1. REGISTER USER AND GET JWT
# ============================================================================
def test_register_user():
    global token
    print_test("1. Register Fresh User and Get JWT")
    try:
        payload = {
            "email": TEST_EMAIL,
            "password": TEST_PASSWORD,
            "name": TEST_NAME
        }
        response = requests.post(f"{BASE_URL}/auth/register", json=payload, timeout=10)
        print_info(f"Status: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            if "token" in data and "user" in data:
                token = data["token"]
                print_success(f"User registered successfully. Email: {TEST_EMAIL}")
                print_info(f"Token: {token[:30]}...")
                return True
            else:
                print_error("Missing token or user in response")
                return False
        else:
            print_error(f"Expected 200, got {response.status_code}: {response.text}")
            return False
    except Exception as e:
        print_error(f"Registration failed: {str(e)}")
        return False

# ============================================================================
# 2. TXT UPLOAD
# ============================================================================
def test_txt_upload():
    print_test("2. TXT Upload - Small Resume (~200 chars)")
    try:
        # Create TXT file
        txt_content = """John Doe
Senior QA Automation Engineer
5+ years experience with Playwright, TypeScript, Selenium
Skills: Playwright, TypeScript, SQL, API Testing, CI/CD
Built comprehensive test automation frameworks and pipelines"""
        
        txt_path = "/tmp/test_resume.txt"
        with open(txt_path, "w") as f:
            f.write(txt_content)
        
        print_info(f"Created TXT file: {txt_path} ({len(txt_content)} chars)")
        
        # Upload
        headers = {"Authorization": f"Bearer {token}"}
        with open(txt_path, "rb") as f:
            files = {"file": ("test_resume.txt", f, "text/plain")}
            response = requests.post(f"{BASE_URL}/profile/resume-upload", 
                                   headers=headers, files=files, timeout=15)
        
        print_info(f"Status: {response.status_code}")
        print_info(f"Response: {response.text[:500]}")
        
        if response.status_code == 200:
            data = response.json()
            user = data.get("user", {})
            chars = data.get("chars", 0)
            filename = data.get("filename", "")
            resume_text = user.get("resume_text", "")
            
            # Validate response
            if filename != "test_resume.txt":
                print_error(f"Filename mismatch: expected 'test_resume.txt', got '{filename}'")
                return False
            
            if chars <= 0:
                print_error(f"Chars should be > 0, got {chars}")
                return False
            
            if not resume_text:
                print_error("resume_text is empty")
                return False
            
            # Check if resume_text contains substring from uploaded content
            if "Playwright" not in resume_text or "QA" not in resume_text:
                print_error(f"resume_text doesn't contain expected content. Got: {resume_text[:200]}")
                return False
            
            print_info(f"Chars extracted: {chars}")
            print_info(f"Filename: {filename}")
            print_info(f"Resume text preview: {resume_text[:100]}...")
            print_success("TXT upload successful with correct data")
            return True
        else:
            print_error(f"Expected 200, got {response.status_code}: {response.text}")
            return False
    except Exception as e:
        print_error(f"TXT upload failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# ============================================================================
# 3. PDF UPLOAD (using pre-generated test PDF)
# ============================================================================
def test_pdf_upload():
    print_test("3. PDF Upload - Using pre-generated /tmp/test_resume.pdf")
    try:
        # Use pre-generated PDF
        pdf_path = "/tmp/test_resume.pdf"
        
        # Check if file exists
        if not os.path.exists(pdf_path):
            print_error(f"Pre-generated PDF not found at {pdf_path}")
            return False
        
        file_size = os.path.getsize(pdf_path)
        print_info(f"Using PDF file: {pdf_path} ({file_size} bytes)")
        
        # Upload
        headers = {"Authorization": f"Bearer {token}"}
        with open(pdf_path, "rb") as f:
            files = {"file": ("test_resume.pdf", f, "application/pdf")}
            response = requests.post(f"{BASE_URL}/profile/resume-upload", 
                                   headers=headers, files=files, timeout=15)
        
        print_info(f"Status: {response.status_code}")
        print_info(f"Response: {response.text[:500]}")
        
        if response.status_code == 200:
            data = response.json()
            user = data.get("user", {})
            chars = data.get("chars", 0)
            filename = data.get("filename", "")
            resume_text = user.get("resume_text", "")
            
            # Validate response
            if filename != "test_resume.pdf":
                print_error(f"Filename mismatch: expected 'test_resume.pdf', got '{filename}'")
                return False
            
            if chars <= 0:
                print_error(f"Chars should be > 0, got {chars}")
                return False
            
            if not resume_text:
                print_error("resume_text is empty")
                return False
            
            # Check if resume_text contains expected content (Playwright or QA)
            if "Playwright" not in resume_text and "QA" not in resume_text:
                print_error(f"resume_text doesn't contain expected content. Got: {resume_text[:200]}")
                return False
            
            print_info(f"✅ CRITICAL: Chars extracted: {chars}")
            print_info(f"✅ CRITICAL: Filename: {filename}")
            print_info(f"✅ CRITICAL: Resume text preview (first 200 chars): {resume_text[:200]}...")
            print_success("PDF upload successful with correct data - unpdf parser working!")
            return True
        else:
            print_error(f"Expected 200, got {response.status_code}: {response.text}")
            return False
    except Exception as e:
        print_error(f"PDF upload failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# ============================================================================
# 4. DOCX UPLOAD (using python-docx)
# ============================================================================
def test_docx_upload():
    print_test("4. DOCX Upload - Generate with python-docx")
    try:
        # Install python-docx if needed
        print_info("Installing python-docx...")
        os.system("pip install python-docx --break-system-packages -q")
        
        from docx import Document
        
        # Create DOCX
        docx_path = "/tmp/test_resume.docx"
        doc = Document()
        
        # Add resume content
        doc.add_paragraph("John Doe")
        doc.add_paragraph("Senior QA Automation Engineer")
        doc.add_paragraph("5+ years experience with Playwright, TypeScript, Selenium, and API Testing. Strong background in building CI/CD test pipelines and automation frameworks.")
        
        doc.save(docx_path)
        print_info(f"Created DOCX file: {docx_path}")
        
        # Upload
        headers = {"Authorization": f"Bearer {token}"}
        with open(docx_path, "rb") as f:
            files = {"file": ("test_resume.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            response = requests.post(f"{BASE_URL}/profile/resume-upload", 
                                   headers=headers, files=files, timeout=15)
        
        print_info(f"Status: {response.status_code}")
        print_info(f"Response: {response.text[:500]}")
        
        if response.status_code == 200:
            data = response.json()
            user = data.get("user", {})
            chars = data.get("chars", 0)
            filename = data.get("filename", "")
            resume_text = user.get("resume_text", "")
            
            # Validate response
            if filename != "test_resume.docx":
                print_error(f"Filename mismatch: expected 'test_resume.docx', got '{filename}'")
                return False
            
            if chars <= 0:
                print_error(f"Chars should be > 0, got {chars}")
                return False
            
            if not resume_text:
                print_error("resume_text is empty")
                return False
            
            # Check if resume_text contains expected content
            if "Playwright" not in resume_text or "QA" not in resume_text:
                print_error(f"resume_text doesn't contain expected content. Got: {resume_text[:200]}")
                return False
            
            print_info(f"Chars extracted: {chars}")
            print_info(f"Filename: {filename}")
            print_info(f"Resume text preview: {resume_text[:150]}...")
            print_success("DOCX upload successful with correct data")
            return True
        else:
            print_error(f"Expected 200, got {response.status_code}: {response.text}")
            return False
    except Exception as e:
        print_error(f"DOCX upload failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# ============================================================================
# 5. MD UPLOAD
# ============================================================================
def test_md_upload():
    print_test("5. MD Upload - Markdown Resume")
    try:
        # Create MD file
        md_content = """# John Doe
## Senior QA Automation Engineer

### Experience
- 5+ years with Playwright, TypeScript, Selenium
- Built comprehensive test automation frameworks
- CI/CD pipeline integration

### Skills
- Playwright
- TypeScript
- SQL
- API Testing
"""
        
        md_path = "/tmp/test_resume.md"
        with open(md_path, "w") as f:
            f.write(md_content)
        
        print_info(f"Created MD file: {md_path} ({len(md_content)} chars)")
        
        # Upload
        headers = {"Authorization": f"Bearer {token}"}
        with open(md_path, "rb") as f:
            files = {"file": ("test_resume.md", f, "text/markdown")}
            response = requests.post(f"{BASE_URL}/profile/resume-upload", 
                                   headers=headers, files=files, timeout=15)
        
        print_info(f"Status: {response.status_code}")
        print_info(f"Response: {response.text[:500]}")
        
        if response.status_code == 200:
            data = response.json()
            user = data.get("user", {})
            chars = data.get("chars", 0)
            filename = data.get("filename", "")
            resume_text = user.get("resume_text", "")
            
            # Validate response
            if filename != "test_resume.md":
                print_error(f"Filename mismatch: expected 'test_resume.md', got '{filename}'")
                return False
            
            if chars <= 0:
                print_error(f"Chars should be > 0, got {chars}")
                return False
            
            if not resume_text:
                print_error("resume_text is empty")
                return False
            
            # Check if resume_text contains substring from uploaded content
            if "Playwright" not in resume_text or "QA" not in resume_text:
                print_error(f"resume_text doesn't contain expected content. Got: {resume_text[:200]}")
                return False
            
            print_info(f"Chars extracted: {chars}")
            print_info(f"Filename: {filename}")
            print_info(f"Resume text preview: {resume_text[:100]}...")
            print_success("MD upload successful with correct data")
            return True
        else:
            print_error(f"Expected 200, got {response.status_code}: {response.text}")
            return False
    except Exception as e:
        print_error(f"MD upload failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# ============================================================================
# 6. UNSUPPORTED FILE TYPE (.jpg) - EXPECT 400
# ============================================================================
def test_unsupported_file():
    print_test("6. Unsupported File Type (.jpg) - Expect 400")
    try:
        # Create a fake JPG file
        jpg_path = "/tmp/test_image.jpg"
        with open(jpg_path, "wb") as f:
            f.write(b"\xFF\xD8\xFF\xE0\x00\x10JFIF")  # JPG header
        
        print_info(f"Created JPG file: {jpg_path}")
        
        # Upload
        headers = {"Authorization": f"Bearer {token}"}
        with open(jpg_path, "rb") as f:
            files = {"file": ("test_image.jpg", f, "image/jpeg")}
            response = requests.post(f"{BASE_URL}/profile/resume-upload", 
                                   headers=headers, files=files, timeout=15)
        
        print_info(f"Status: {response.status_code}")
        print_info(f"Response: {response.text}")
        
        if response.status_code == 400:
            data = response.json()
            error = data.get("error", "")
            if "Unsupported file type" in error and ".jpg" in error:
                print_success("Unsupported file type correctly rejected with 400")
                return True
            else:
                print_error(f"Expected error message about unsupported file type, got: {error}")
                return False
        elif response.status_code == 500:
            print_error("CRITICAL: Got 500 instead of 400 for unsupported file type")
            return False
        else:
            print_error(f"Expected 400, got {response.status_code}")
            return False
    except Exception as e:
        print_error(f"Unsupported file test failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# ============================================================================
# 7. MISSING FILE FIELD - EXPECT 400
# ============================================================================
def test_missing_file():
    print_test("7. Missing File Field - Expect 400")
    try:
        # Upload without file field
        headers = {"Authorization": f"Bearer {token}"}
        response = requests.post(f"{BASE_URL}/profile/resume-upload", 
                               headers=headers, files={}, timeout=15)
        
        print_info(f"Status: {response.status_code}")
        print_info(f"Response: {response.text}")
        
        if response.status_code == 400:
            data = response.json()
            error = data.get("error", "")
            if "No file uploaded" in error:
                print_success("Missing file correctly rejected with 400")
                return True
            else:
                print_error(f"Expected error message about no file uploaded, got: {error}")
                return False
        elif response.status_code == 500:
            print_error("CRITICAL: Got 500 instead of 400 for missing file")
            return False
        else:
            print_error(f"Expected 400, got {response.status_code}")
            return False
    except Exception as e:
        print_error(f"Missing file test failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# ============================================================================
# 8. MISSING AUTH HEADER - EXPECT 401
# ============================================================================
def test_missing_auth():
    print_test("8. Missing Auth Header - Expect 401")
    try:
        # Create a dummy file
        txt_path = "/tmp/dummy.txt"
        with open(txt_path, "w") as f:
            f.write("dummy content")
        
        # Upload without auth header
        with open(txt_path, "rb") as f:
            files = {"file": ("dummy.txt", f, "text/plain")}
            response = requests.post(f"{BASE_URL}/profile/resume-upload", 
                                   files=files, timeout=15)
        
        print_info(f"Status: {response.status_code}")
        print_info(f"Response: {response.text}")
        
        if response.status_code == 401:
            print_success("Missing auth correctly rejected with 401")
            return True
        else:
            print_error(f"Expected 401, got {response.status_code}")
            return False
    except Exception as e:
        print_error(f"Missing auth test failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# ============================================================================
# 9. CORRUPT PDF - EXPECT 400 (NOT 500)
# ============================================================================
def test_corrupt_pdf():
    print_test("9. Corrupt PDF - Expect 400 (NOT 500)")
    try:
        # Create corrupt PDF with random bytes
        corrupt_path = "/tmp/corrupt.pdf"
        with open(corrupt_path, "wb") as f:
            f.write(b"This is not a valid PDF file, just random bytes: \x00\x01\x02\x03\xFF\xFE")
        
        print_info(f"Created corrupt PDF: {corrupt_path}")
        
        # Upload
        headers = {"Authorization": f"Bearer {token}"}
        with open(corrupt_path, "rb") as f:
            files = {"file": ("corrupt.pdf", f, "application/pdf")}
            response = requests.post(f"{BASE_URL}/profile/resume-upload", 
                                   headers=headers, files=files, timeout=15)
        
        print_info(f"Status: {response.status_code}")
        print_info(f"Response: {response.text}")
        
        if response.status_code == 400:
            data = response.json()
            error = data.get("error", "")
            if "Could not parse PDF" in error or "parse" in error.lower():
                print_success("Corrupt PDF correctly rejected with 400 and parser error message")
                return True
            else:
                print_info(f"Got 400 but with different error message: {error}")
                print_success("Corrupt PDF rejected with 400 (acceptable)")
                return True
        elif response.status_code == 500:
            print_error("CRITICAL: Got 500 instead of 400 for corrupt PDF")
            return False
        else:
            print_error(f"Expected 400, got {response.status_code}")
            return False
    except Exception as e:
        print_error(f"Corrupt PDF test failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# ============================================================================
# 10. VERIFY GET /api/auth/me SHOWS UPDATED RESUME DATA
# ============================================================================
def test_auth_me_resume_data():
    print_test("10. Verify GET /api/auth/me Shows Updated Resume Data")
    try:
        headers = {"Authorization": f"Bearer {token}"}
        response = requests.get(f"{BASE_URL}/auth/me", headers=headers, timeout=10)
        
        print_info(f"Status: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            user = data.get("user", {})
            resume_text = user.get("resume_text", "")
            resume_filename = user.get("resume_filename", "")
            
            if not resume_text:
                print_error("resume_text is empty in /auth/me response")
                return False
            
            if not resume_filename:
                print_error("resume_filename is empty in /auth/me response")
                return False
            
            print_info(f"Resume filename: {resume_filename}")
            print_info(f"Resume text length: {len(resume_text)} chars")
            print_info(f"Resume text preview: {resume_text[:100]}...")
            print_success("GET /api/auth/me shows updated resume data")
            return True
        else:
            print_error(f"Expected 200, got {response.status_code}: {response.text}")
            return False
    except Exception as e:
        print_error(f"Auth me resume data test failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# ============================================================================
# 11. SANITY CHECK - GET /api/dashboard/stats
# ============================================================================
def test_dashboard_stats_sanity():
    print_test("11. Sanity Check - GET /api/dashboard/stats")
    try:
        headers = {"Authorization": f"Bearer {token}"}
        response = requests.get(f"{BASE_URL}/dashboard/stats", headers=headers, timeout=10)
        
        print_info(f"Status: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            stats = data.get("stats", {})
            resume_score = data.get("resume_score")
            
            print_info(f"Stats: {stats}")
            print_info(f"Resume score: {resume_score}")
            print_success("GET /api/dashboard/stats working correctly")
            return True
        else:
            print_error(f"Expected 200, got {response.status_code}: {response.text}")
            return False
    except Exception as e:
        print_error(f"Dashboard stats sanity check failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# ============================================================================
# 12. SANITY CHECK - GET /api/jobs
# ============================================================================
def test_jobs_sanity():
    print_test("12. Sanity Check - GET /api/jobs")
    try:
        response = requests.get(f"{BASE_URL}/jobs", timeout=10)
        
        print_info(f"Status: {response.status_code}")
        
        if response.status_code == 200:
            data = response.json()
            jobs = data.get("jobs", [])
            print_info(f"Found {len(jobs)} jobs")
            print_success("GET /api/jobs working correctly")
            return True
        else:
            print_error(f"Expected 200, got {response.status_code}: {response.text}")
            return False
    except Exception as e:
        print_error(f"Jobs sanity check failed: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

# ============================================================================
# MAIN TEST RUNNER
# ============================================================================
def main():
    print("\n" + "="*80)
    print("JobOS AI Phase 3 - Resume Upload Endpoint Test Suite")
    print("Testing FIXED pdf-parse v2 PDFParse class implementation")
    print("="*80)
    print(f"Base URL: {BASE_URL}")
    print(f"Test Email: {TEST_EMAIL}")
    print("="*80)
    
    results = {}
    
    # Run all tests in sequence
    tests = [
        ("1. Register User", test_register_user),
        ("2. TXT Upload", test_txt_upload),
        ("3. PDF Upload", test_pdf_upload),
        ("4. DOCX Upload", test_docx_upload),
        ("5. MD Upload", test_md_upload),
        ("6. Unsupported File (.jpg)", test_unsupported_file),
        ("7. Missing File Field", test_missing_file),
        ("8. Missing Auth Header", test_missing_auth),
        ("9. Corrupt PDF", test_corrupt_pdf),
        ("10. Auth Me Resume Data", test_auth_me_resume_data),
        ("11. Dashboard Stats Sanity", test_dashboard_stats_sanity),
        ("12. Jobs Sanity", test_jobs_sanity),
    ]
    
    for test_name, test_func in tests:
        try:
            results[test_name] = test_func()
        except Exception as e:
            print_error(f"Test '{test_name}' crashed: {str(e)}")
            import traceback
            traceback.print_exc()
            results[test_name] = False
        
        # Small delay between tests
        time.sleep(0.5)
    
    # Print summary
    print("\n" + "="*80)
    print("TEST SUMMARY")
    print("="*80)
    
    passed = sum(1 for v in results.values() if v)
    failed = sum(1 for v in results.values() if not v)
    
    for test_name, result in results.items():
        status = "✅ PASS" if result else "❌ FAIL"
        print(f"{status}: {test_name}")
    
    print("="*80)
    print(f"Total: {len(results)} | Passed: {passed} | Failed: {failed}")
    print("="*80)
    
    # File type summary
    print("\n" + "="*80)
    print("FILE TYPE PARSING SUMMARY")
    print("="*80)
    file_types = {
        "TXT": results.get("2. TXT Upload", False),
        "PDF": results.get("3. PDF Upload", False),
        "DOCX": results.get("4. DOCX Upload", False),
        "MD": results.get("5. MD Upload", False),
    }
    
    for file_type, success in file_types.items():
        status = "✅ Successfully parsed" if success else "❌ Failed to parse"
        print(f"{status}: {file_type}")
    
    print("="*80)
    
    # Check for 500 errors
    print("\n" + "="*80)
    print("500 ERROR CHECK")
    print("="*80)
    if all([results.get("6. Unsupported File (.jpg)", False),
            results.get("7. Missing File Field", False),
            results.get("9. Corrupt PDF", False)]):
        print("✅ NO 500 ERRORS - All error cases returned 400/401 as expected")
    else:
        print("❌ POTENTIAL 500 ERRORS - Check test output above")
    print("="*80)
    
    if failed == 0:
        print("\n🎉 ALL TESTS PASSED!")
        return 0
    else:
        print(f"\n⚠️  {failed} TEST(S) FAILED")
        return 1

if __name__ == "__main__":
    exit(main())
